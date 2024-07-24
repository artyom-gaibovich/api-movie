import os
from docx import Document

def read_directory(directory):
    document = Document()N
    for root, _, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            document.add_heading(file, level=1)
            document.add_paragraph(content)
            document.add_paragraph()

    return document

if __name__ == "__main__":
    src_directory = 'src'
    document = read_directory(src_directory)
    document.save('report.docx')
